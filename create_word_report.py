import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import pandas as pd
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from oauth2client.service_account import ServiceAccountCredentials
import gspread
import os

# Function to download the image from a URL
def download_image(service_account_file, file_path, save_path):
    # Authenticate using service account credentials
    credentials = service_account.Credentials.from_service_account_file(service_account_file)
    drive_service = build('drive', 'v3', credentials=credentials, cache_discovery=False)
    file_id = get_file_id_from_url(file_path)
    # Download the image
    request = drive_service.files().get_media(fileId=file_id)
    fh = open(save_path, "wb")
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while done is False:
        status, done = downloader.next_chunk()
        print("Download %d%%." % int(status.progress() * 100))
    print("Image downloaded successfully.")

def get_file_id_from_url(url):
    # Extract the file ID from the URL
    return url.split('/')[-2]

############################ Google Sheets ##########################################
import gspread

# Define the scope and credentials for Google Sheets API
scope = ['https://spreadsheets.google.com/feeds',
   'https://www.googleapis.com/auth/drive']

# Authorize the client with credentials
creds = ServiceAccountCredentials.from_json_keyfile_name('/home/sailbslsafety/mysite/credentials.json', scope)
client = gspread.authorize(creds)

#########################################




def create_report(observations_df):
    # Find the start and end date of observations in the dataframe
    start_date = observations_df['date'].min()
    end_date = observations_df['date'].max()

    #Find the department's name of which there are observations in the dataframe
    department = observations_df.loc[0, 'department']
    doc = docx.Document()

    para=doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r=para.add_run(str(department))
    r.bold = True

    para=doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r=para.add_run('Inspection Observations')
    r.bold = True

    para=doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r=para.add_run(str(start_date) + " to " + str(end_date))
    r.bold = True

    para=doc.add_paragraph()

    table = doc.add_table(rows=observations_df.shape[0]+1, cols=4)
    table.style = 'Table Grid'
    row = table.rows[0].cells
    row[0].text = "Location"
    row[0].paragraphs[0].runs[0].font.bold = True
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row[1].text = "Observation"
    row[1].paragraphs[0].runs[0].font.bold = True
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row[2].text = "Picture 1"
    row[2].paragraphs[0].runs[0].font.bold = True
    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row[3].text = "Picture 2"
    row[3].paragraphs[0].runs[0].font.bold = True
    row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    i=1
    for index, observation in observations_df.iterrows():
        row = table.rows[i].cells
        row[0].text = observation['location']
        row[1].text = observation['observation']

        if (observation['photo']==''):
            pass
        elif (len(observation['photo'].split(","))==1):
            paragraph = row[2].paragraphs[0]
            image_1 = observation['photo'].split(",")[0]  #get a list of links to all images of the particular observation
            download_image('/home/sailbslsafety/mysite/credentials.json', image_1, 'image.jpg')
            paragraph.add_run().add_picture('image.jpg', width=Inches(2.0), height=Inches(2.0))
            # Remove the image from local system
            os.remove('image.jpg')
        elif (len(observation['photo'].split(","))>1):
            paragraph = row[2].paragraphs[0]
            image_1 = observation['photo'].split(",")[0]  #get a list of links to all images of the particular observation
            download_image('/home/sailbslsafety/mysite/credentials.json', image_1, 'image.jpg')
            paragraph.add_run().add_picture('image.jpg', width=Inches(2.0), height=Inches(2.0))
            # Remove the image from local system
            os.remove('image.jpg')

            paragraph = row[3].paragraphs[0]
            image_2 = observation['photo'].split(",")[1]  #get a list of links to all images of the particular observation
            download_image('/home/sailbslsafety/mysite/credentials.json', image_2, 'image.jpg')
            paragraph.add_run().add_picture('image.jpg', width=Inches(2.0), height=Inches(2.0))
            # Remove the image from local system
            os.remove('image.jpg')
        else:
            pass
        i+=1
    file_path = "/home/sailbslsafety/mysite/downloads" + str(department) + "-" + str(start_date) + " to " + str(end_date) +".docx"
    doc.save(file_path)
    return file_path
